import io
from docx import Document

# function to extract data from the document
def extract_text_from_doc(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        text =[]
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
                        
        return "\n".join(text)
    except Exception as e:
        print(f"Error reading document: {e}")
        return ""